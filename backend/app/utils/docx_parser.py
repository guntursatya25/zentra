"""DOCX parser using python-docx — extracts text with structure detection."""

import re
from typing import Any

from docx import Document as DocxDocument


class DOCXParser:
    """Parse DOCX documents, extracting text with Bab/Pasal/Ayat structure."""

    BAB_PATTERN = re.compile(
        r"^BAB\s+([IVXLCDM]+|[0-9]+|[A-Z])\b",
        re.IGNORECASE,
    )
    PASAL_PATTERN = re.compile(
        r"^Pasal\s+([0-9]+[A-Za-z]?|[A-Z])\b",
        re.IGNORECASE,
    )
    AYAT_PATTERN = re.compile(
        r"^Ayat\s*\((\d+[a-z]?)\)\b",
        re.IGNORECASE,
    )

    def __init__(self, file_path: str):
        self.file_path = file_path

    def extract_sections(self) -> list[dict[str, Any]]:
        """Extract structured sections from DOCX.

        Uses heading styles when available, falls back to regex pattern matching
        on paragraph text.

        Returns list of dicts with keys:
            teks, bab, bab_judul, pasal, pasal_judul, ayat, halaman
        """
        doc = DocxDocument(self.file_path)
        sections: list[dict[str, Any]] = []
        current_bab = None
        current_bab_judul = None
        current_pasal = None
        current_pasal_judul = None
        current_ayat = None
        buffer: list[str] = []
        halaman = 1  # GAP 13: Start at page 1

        def flush_buffer() -> None:
            if not buffer:
                return
            teks = " ".join(buffer).strip()
            if teks and len(teks) > 5:  # Skip very short fragments
                sections.append({
                    "teks": teks,
                    "bab": current_bab,
                    "bab_judul": current_bab_judul,
                    "pasal": current_pasal,
                    "pasal_judul": current_pasal_judul,
                    "ayat": current_ayat,
                    "halaman": halaman,
                })

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name.lower() if para.style else ""
            is_heading = style_name.startswith("heading")

            # Try heading style first, then regex
            bab_match = self.BAB_PATTERN.match(text)
            pasal_match = self.PASAL_PATTERN.match(text)
            ayat_match = self.AYAT_PATTERN.match(text)

            if bab_match and (is_heading or len(text) < 100):
                flush_buffer()
                current_bab = bab_match.group(1)
                current_bab_judul = None
                current_pasal = None
                current_pasal_judul = None
                current_ayat = None
                buffer = [text]
                continue

            if pasal_match and (is_heading or len(text) < 100):
                flush_buffer()
                current_pasal = pasal_match.group(1)
                current_pasal_judul = None
                current_ayat = None
                buffer = [text]
                continue

            if ayat_match:
                flush_buffer()
                current_ayat = ayat_match.group(1)
                buffer = [text]
                continue

            # Track page breaks (Word manual page breaks)
            if text == "\f" or "PAGE BREAK" in text.upper():
                halaman += 1  # GAP 13: Increment page counter
                continue

            # Track judgment lines as titles
            if current_bab and current_bab_judul is None and len(buffer) <= 1:
                current_bab_judul = text
                buffer.append(text)
                continue

            if current_pasal and current_pasal_judul is None and len(buffer) <= 1:
                current_pasal_judul = text
                buffer.append(text)
                continue

            buffer.append(text)

        flush_buffer()
        return sections

    def extract_title(self) -> str | None:
        """Extract document title from first heading or paragraph."""
        doc = DocxDocument(self.file_path)
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name.lower() if para.style else ""
            if style_name.startswith("heading") or style_name == "title":
                return text

        # Fallback: first substantial paragraph
        for para in doc.paragraphs:
            text = para.text.strip()
            if len(text) > 10:
                return text

        return None
